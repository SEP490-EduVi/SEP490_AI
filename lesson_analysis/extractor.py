"""Extract text content from PDF and DOCX files."""

import os
import logging
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_pdf(file_path: str) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    text_parts: list[str] = []

    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)
        logger.info("PDF has %d page(s)", total_pages)

        for i, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i} ---\n{page_text}")
            else:
                text_parts.append(f"--- Page {i} ---\n[No extractable text]")

    return "\n\n".join(text_parts)


def extract_docx(file_path: str) -> str:
    """Extract all text from a DOCX file using python-docx."""
    doc = Document(file_path)
    text_parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table_idx, table in enumerate(doc.tables, start=1):
        table_rows: list[str] = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_rows.append(row_text)
        if table_rows:
            text_parts.append(f"\n--- Table {table_idx} ---")
            text_parts.extend(table_rows)

    logger.info(
        "DOCX has %d paragraph(s) and %d table(s)",
        len(doc.paragraphs),
        len(doc.tables),
    )

    return "\n".join(text_parts)


def extract_text(file_path: str) -> str:
    """
    Detect file type and delegate to the appropriate extractor.

    Raises:
        ValueError: If the file type is not supported.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. Only .pdf and .docx are supported."
        )
