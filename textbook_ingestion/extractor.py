"""Extract text from PDF (Vision OCR) and DOCX files."""

import json
import os
import re
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from google.cloud import vision

logger = logging.getLogger(__name__)

_VISION_BATCH_SIZE = 5  # Max pages per synchronous Vision API batch

_vision_client: Optional[vision.ImageAnnotatorClient] = None


def _get_vision_client() -> vision.ImageAnnotatorClient:
    """Lazy-init and return the Vision API client."""
    global _vision_client
    if _vision_client is None:
        _vision_client = vision.ImageAnnotatorClient()
        logger.info("Google Cloud Vision client initialised.")
    return _vision_client


# ── OCR cache ──


def _cache_dir_for(file_path: str) -> Path:
    return Path(file_path).with_suffix(".ocr_cache")


def _cache_path_for_page(file_path: str, page_num: int) -> Path:
    return _cache_dir_for(file_path) / f"page_{page_num:04d}.json"


def _read_cached_page(file_path: str, page_num: int) -> Optional[str]:
    """Return cached OCR text for a page, or None on miss."""
    cache_file = _cache_path_for_page(file_path, page_num)
    if cache_file.exists():
        try:
            data = json.loads(cache_file.read_text(encoding="utf-8"))
            return data.get("text")
        except (json.JSONDecodeError, KeyError):
            return None
    return None


def _write_cached_page(file_path: str, page_num: int, text: str) -> None:
    cache_dir = _cache_dir_for(file_path)
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = _cache_path_for_page(file_path, page_num)
    cache_file.write_text(
        json.dumps({"page": page_num, "text": text}, ensure_ascii=False),
        encoding="utf-8",
    )


# ── PDF page count ──


def _count_pdf_pages(file_path: str) -> int:
    """Count pages by scanning for /Type /Page entries (no external tools)."""
    with open(file_path, "rb") as f:
        data = f.read()

    pages = re.findall(rb"/Type\s*/Page(?!s)", data)
    count = len(pages)

    if count == 0:
        raise ValueError(
            f"Could not determine page count for {file_path}. "
            "The file may be corrupted or not a valid PDF."
        )
    logger.debug("Detected %d pages in %s (byte-scan).", count, file_path)
    return count


# ── Vision OCR (native PDF batch) ──


def _ocr_pdf_pages(
    pdf_content: bytes,
    pages: list[int],
) -> dict[int, str]:
    """OCR a batch of PDF pages via Vision batch_annotate_files. Returns {page_num: text}."""
    client = _get_vision_client()

    input_config = vision.InputConfig(
        content=pdf_content,
        mime_type="application/pdf",
    )

    feature = vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)

    request = vision.AnnotateFileRequest(
        input_config=input_config,
        features=[feature],
        pages=pages,
        image_context=vision.ImageContext(language_hints=["vi"]),
    )

    response = client.batch_annotate_files(requests=[request])

    results: dict[int, str] = {}

    for file_response in response.responses:
        if file_response.error.message:
            raise RuntimeError(
                f"Vision API batch error: {file_response.error.message}"
            )
        for i, page_response in enumerate(file_response.responses):
            page_num = pages[i] if i < len(pages) else i + 1
            annotation = page_response.full_text_annotation
            results[page_num] = annotation.text if annotation else ""

    return results


# ── PDF extraction (Vision OCR only, with cache) ─────────────────────


def extract_pdf(file_path: str, *, skip_start: int = 0, skip_end: int = 0) -> str:
    """
    Extract text from every page of a PDF via Google Cloud Vision OCR.

    The PDF is sent directly to Vision API's ``batch_annotate_files``
    endpoint (up to 5 pages per request). No image conversion or Poppler
    needed.

    Cached pages are read from ``<file>.ocr_cache/`` and skipped.
    Newly OCR'd pages are written to the cache so re-runs skip them.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Full text with ``--- Page N ---`` markers.

    Raises:
        RuntimeError: If Vision API returns an error.
        ValueError:   If the PDF is empty or unreadable.
    """
    total_pages = _count_pdf_pages(file_path)
    logger.info("PDF has %d page(s). Starting Vision OCR...", total_pages)

    # Skip leading/trailing pages based on caller arguments
    start_page = skip_start + 1
    end_page = total_pages - skip_end
    if skip_start > 0:
        logger.info("  Skipping first %d page(s) (skip_start=%d).", skip_start, skip_start)
    if skip_end > 0:
        logger.info("  Skipping last %d page(s) (skip_end=%d, last extracted=%d).", skip_end, skip_end, end_page)
    if start_page > end_page:
        raise ValueError(
            f"skip_start={skip_start} + skip_end={skip_end} >= total pages ({total_pages}). Nothing to extract."
        )

    pdf_content = Path(file_path).read_bytes()

    page_texts: dict[int, str] = {}
    pages_to_ocr: list[int] = []
    cached_count = 0

    for page_num in range(start_page, end_page + 1):
        cached_text = _read_cached_page(file_path, page_num)
        if cached_text is not None:
            page_texts[page_num] = cached_text
            cached_count += 1
        else:
            pages_to_ocr.append(page_num)

    logger.info(
        "  %d page(s) cached, %d page(s) to OCR.",
        cached_count, len(pages_to_ocr),
    )

    # OCR uncached pages in batches
    ocr_count = 0
    failed_pages: list[int] = []

    for batch_start in range(0, len(pages_to_ocr), _VISION_BATCH_SIZE):
        batch = pages_to_ocr[batch_start:batch_start + _VISION_BATCH_SIZE]
        logger.info(
            "  OCR batch: pages %s (%d/%d done)",
            batch, ocr_count, len(pages_to_ocr),
        )

        try:
            batch_results = _ocr_pdf_pages(pdf_content, batch)
        except Exception as e:
            logger.error("  Batch OCR failed for pages %s: %s", batch, e)
            for p in batch:
                page_texts[p] = f"[ERROR: Vision OCR failed – {e}]"
                failed_pages.append(p)
            continue

        for page_num, text in batch_results.items():
            page_texts[page_num] = text
            _write_cached_page(file_path, page_num, text)
            ocr_count += 1

            if text.strip():
                logger.info(
                    "  [%d/%d] OK (%d chars)", page_num, total_pages, len(text)
                )
            else:
                logger.warning(
                    "  [%d/%d] OCR returned empty text", page_num, total_pages
                )

    logger.info(
        "PDF extraction complete: %d cached, %d newly OCR'd, %d failed "
        "out of %d pages.",
        cached_count, ocr_count, len(failed_pages), total_pages,
    )
    if failed_pages:
        logger.warning("Failed pages: %s", failed_pages)

    # Assemble in page order (respecting skip_start/skip_end)
    text_parts: list[str] = []
    for page_num in range(start_page, end_page + 1):
        text = page_texts.get(page_num, "")
        text_parts.append(f"--- Page {page_num} ---\n{text}")

    return "\n\n".join(text_parts)


# ── DOCX extraction ──


def extract_docx(file_path: str) -> str:
    """Extract all text and tables from a DOCX file."""
    doc = Document(file_path)
    text_parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table_idx, table in enumerate(doc.tables, start=1):
        table_rows: list[str] = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_rows.append(row_text)
        if table_rows:
            text_parts.append(f"\n--- Table {table_idx} ---")
            text_parts.extend(table_rows)

    logger.info(
        "DOCX has %d paragraph(s) and %d table(s)",
        len(doc.paragraphs),
        len(doc.tables),
    )

    return "\n".join(text_parts)


# ── Entry point ──


def extract_text(file_path: str, *, skip_start: int = 0, skip_end: int = 0) -> str:
    """Route to PDF or DOCX extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_pdf(file_path, skip_start=skip_start, skip_end=skip_end)
    elif ext in (".docx", ".doc"):
        return extract_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. Only .pdf and .docx are supported."
        )
